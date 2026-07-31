"""DocumentParser — 解析 PDF/DOCX 简历文件，提取文本和图片

纯代码工具，不调用 LLM。被 Skill 2 (ResumeExtractor) 作为预处理步骤调用。

复用自参考仓库 ResumeScreening_ref/src/parsers/document_parser.py，
适配 F1 项目的异常体系和配置。
"""

from pathlib import Path

from docx import Document
import fitz

from backend.core.exceptions import ParseException
from backend.core.logger import get_logger

SUPPORTED_FORMATS: frozenset[str] = frozenset({".pdf", ".docx", ".doc"})
MIN_IMAGE_SIZE = 1024  # 最小图片字节数（过滤图标）

# 常见简历章节标题 —— 用于插入结构标记，保留文档骨架
_SECTION_PATTERNS = [
    "基本信息", "个人概况", "个人信息", "求职意向", "自我评价", "个人评价",
    "工作经历", "工作经验", "工作履历", "职业经历", "从业经验",
    "项目经历", "项目经验", "主要项目", "代表项目", "重点项目",
    "教育背景", "教育经历", "学习经历", "学历",
    "技能", "专业技能", "技术栈", "技术能力", "掌握技能",
    "证书", "资质证书", "语言能力", "获奖", "荣誉",
    "实习经历", "社会实践", "校园经历", "社团", "志愿者",
    "培训经历", "研究经历", "论文", "专利",
]


class DocumentParser:
    """简历文档解析器，支持 PDF / DOCX"""

    def __init__(self):
        self._log = get_logger()

    def parse(self, file_path: str | Path) -> tuple[str, list[bytes]]:
        """解析文档，同时提取文本和图片

        Args:
            file_path: 文件路径

        Returns:
            (text, images): 纯文本 + 图片字节列表

        Raises:
            ParseException: 文件不存在/格式不支持/解析失败
        """
        file_path = Path(file_path)
        self._validate(file_path)
        suffix = file_path.suffix.lower()

        self._log.info(f"开始解析文档: {file_path.name} ({suffix})")

        if suffix == ".pdf":
            text, images = self._parse_pdf(file_path)
        else:
            text, images = self._parse_docx(file_path)

        self._log.info(
            f"文档解析完成: {file_path.name} | "
            f"chars={len(text)}, images={len(images)}"
        )
        return text, images

    def parse_text_only(self, file_path: str | Path) -> str:
        """只提取文本"""
        text, _ = self.parse(file_path)
        return text

    # ── 内部方法 ──────────────────────────────────────────

    def _validate(self, file_path: Path) -> None:
        if not file_path.exists():
            raise ParseException(
                message=f"文件不存在: {file_path}",
                file_name=file_path.name,
            )
        if not file_path.is_file():
            raise ParseException(
                message=f"路径不是文件: {file_path}",
                file_name=file_path.name,
            )
        if file_path.suffix.lower() not in SUPPORTED_FORMATS:
            raise ParseException(
                message=f"不支持的文件格式: {file_path.suffix}，支持: {SUPPORTED_FORMATS}",
                file_type=file_path.suffix,
                file_name=file_path.name,
            )

    def _parse_pdf(self, file_path: Path) -> tuple[str, list[bytes]]:
        text_parts: list[str] = []
        images: list[bytes] = []
        seen_xrefs: set[int] = set()

        try:
            doc = fitz.open(file_path)
            for page_idx in range(doc.page_count):
                page = doc[page_idx]
                page_text = page.get_text()
                if page_text and isinstance(page_text, str) and page_text.strip():
                    text_parts.append(page_text.strip())

                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    if xref in seen_xrefs:
                        continue
                    seen_xrefs.add(xref)
                    try:
                        base_image = doc.extract_image(xref)
                        img_bytes = base_image["image"]
                        if len(img_bytes) > MIN_IMAGE_SIZE:
                            images.append(img_bytes)
                    except Exception:
                        pass  # 单张图片提取失败不影响整体
            doc.close()
        except Exception as e:
            raise ParseException(
                message=f"PDF 解析失败: {e}",
                file_type="pdf",
                file_name=file_path.name,
                details={"error": str(e)},
            ) from e

        raw = "\n\n".join(text_parts)
        return self._add_section_markers(raw), images

    def _parse_docx(self, file_path: Path) -> tuple[str, list[bytes]]:
        text_parts: list[str] = []
        images: list[bytes] = []

        try:
            doc = Document(str(file_path))

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        img_bytes = rel.target_part.blob
                        if len(img_bytes) > MIN_IMAGE_SIZE:
                            images.append(img_bytes)
                    except Exception:
                        pass

        except Exception as e:
            raise ParseException(
                message=f"DOCX 解析失败: {e}",
                file_type="docx",
                file_name=file_path.name,
                details={"error": str(e)},
            ) from e

        return "\n\n".join(text_parts), images

    def _add_section_markers(self, text: str) -> str:
        """在识别到的简历章节标题前插入 [SECTION: 名称] 标记，保留文档骨架。"""
        lines = text.split("\n")
        result = []
        for line in lines:
            stripped = line.strip()
            # 检查是否为已知章节标题（短行 + 匹配已知模式）
            if len(stripped) <= 20:
                for pattern in _SECTION_PATTERNS:
                    if stripped == pattern or stripped.startswith(pattern):
                        result.append(f"[SECTION: {pattern}]")
                        break
            result.append(line)
        return "\n".join(result)


# 全局单例
document_parser = DocumentParser()
